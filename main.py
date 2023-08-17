# import undetected_chromedriver as uc
from bs4 import BeautifulSoup
import requests
import re
import pandas as pd

import docx
import io

# testing out selenium
# driver = uc.Chrome()
# driver.get('https://e-seimas.lrs.lt')


def get_data(id):

    # search
    url = "https://www.e-tar.lt/portal/lt/legalActSearch"
    data = {
        "javax.faces.partial.ajax": "false",
        "javax.faces.source": "contentForm:searchParamPane:searchButton",
        "javax.faces.partial.execute": "@all",
        "javax.faces.partial.render": "contentForm:resultsPanel contentForm:searchParamPane",
        "contentForm:searchParamPane:searchButton": "contentForm:searchParamPane:searchButton",
        "contentForm_SUBMIT": "1",

        # we can copy this from the end of a failed request
        "javax.faces.ViewState": "rO0ABXVyABNbTGphdmEubGFuZy5PYmplY3Q7kM5YnxBzKWwCAAB4cAAAAAJ1cQB+AAAAAAACcHNyABFqYXZhLnV0aWwuSGFzaE1hcAUH2sHDFmDRAwACRgAKbG9hZEZhY3RvckkACXRocmVzaG9sZHhwP0AAAAAAADB3CAAAAEAAAAAbdAAsY29udGVudEZvcm06c2VhcmNoUGFyYW1QYW5lOmpfaWRfNTQ6Y2FsZW5kYXJ1cQB+AAAAAAACdXEAfgAAAAAAAXVxAH4AAAAAAARwdXEAfgAAAAAABH5yAC5qYXZheC5mYWNlcy5jb21wb25lbnQuVUlDb21wb25lbnQkUHJvcGVydHlLZXlzAAAAAAAAAAASAAB4cgAOamF2YS5sYW5nLkVudW0AAAAAAAAAABIAAHhwdAAIYmluZGluZ3NzcgAramF2YXguZmFjZXMuY29tcG9uZW50Ll9BdHRhY2hlZFN0YXRlV3JhcHBlckSr5kB900/EAgACTAAGX2NsYXNzdAARTGphdmEvbGFuZy9DbGFzcztMABNfd3JhcHBlZFN0YXRlT2JqZWN0dAASTGphdmEvbGFuZy9PYmplY3Q7eHB2cgAzamF2YXguZmFjZXMuY29tcG9uZW50Ll9EZWx0YVN0YXRlSGVscGVyJEludGVybmFsTWFwhHIXGwejCVsCAAB4cQB+AAN1cQB+AAAAAAACdAAFbGFiZWxzcgA+b3JnLmFwYWNoZS5teWZhY2VzLnZpZXcuZmFjZWxldHMuZWwuTG9jYXRpb25WYWx1ZUV4cHJlc3Npb25VRUwZUz2DuoEIYAwAAHhyADtvcmcuYXBhY2hlLm15ZmFjZXMudmlldy5mYWNlbGV0cy5lbC5Mb2NhdGlvblZhbHVlRXhwcmVzc2lvbrHF44ghqQEwDAAAeHIAGGphdmF4LmVsLlZhbHVlRXhwcmVzc2lvbncKgFvgwP6RAgAAeHIAE2phdmF4LmVsLkV4cHJlc3Npb26jhYpT8lrSPAIAAHhwc3IARW9yZy5hcGFjaGUubXlmYWNlcy52aWV3LmZhY2VsZXRzLmVsLkNvbnRleHRBd2FyZVRhZ1ZhbHVlRXhwcmVzc2lvblVFTKPhD9oDeAHIDAAAeHIAQm9yZy5hcGFjaGUubXlmYWNlcy52aWV3LmZhY2VsZXRzLmVsLkNvbnRleHRBd2FyZVRhZ1ZhbHVlRXhwcmVzc2lvbgAAAAAAAAABDAAAeHEAfgAYc3IAL29yZy5hcGFjaGUud2ViYmVhbnMuZWwyMi5XcmFwcGVkVmFsdWVFeHByZXNzaW9uAAAAAAAAAAECAAFMAA92YWx1ZUV4cHJlc3Npb250ABpMamF2YXgvZWwvVmFsdWVFeHByZXNzaW9uO3hxAH4AGHNyACFvcmcuYXBhY2hlLmVsLlZhbHVlRXhwcmVzc2lvbkltcGwIjSL+h4ituAwAAHhxAH4AGHclABEje2NjLmF0dHJzLmxhYmVsfQAQamF2YS5sYW5nLk9iamVjdHBweHNyABlqYXZheC5mYWNlcy52aWV3LkxvY2F0aW9uAAAAAAAAAAECAANJAAZjb2x1bW5JAARsaW5lTAAEcGF0aHQAEkxqYXZhL2xhbmcvU3RyaW5nO3hwAAAAhAAAABR0ACQvcmVzb3VyY2VzL2NvbXBvc2l0ZXMvY2FsZW5kYXIueGh0bWx3BwAFdmFsdWV4cQB+ACV3BAAAAAF4fnEAfgAKdAANYXR0cmlidXRlc01hcHNxAH4ADnEAfgATdXEAfgAAAAAAAnQAHG9hbS5DT01NT05fUFJPUEVSVElFU19NQVJLRURzcgAOamF2YS5sYW5nLkxvbmc7i+SQzI8j3wIAAUoABXZhbHVleHIAEGphdmEubGFuZy5OdW1iZXKGrJUdC5TgiwIAAHhwAAAAQAIAAEFwcHB0ACxjb250ZW50Rm9ybTpzZWFyY2hQYXJhbVBhbmU6al9pZF80cTpjYWxlbmRhcnVxAH4AAAAAAAJ1cQB+AAAAAAABdXEAfgAAAAAABHB1cQB+AAAAAAAEcQB+AAxzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AFXNxAH4AFnNxAH4AG3NxAH4AHnNxAH4AIXclABEje2NjLmF0dHJzLmxhYmVsfQAQamF2YS5sYW5nLk9iamVjdHBweHEAfgAldwcABXZhbHVleHEAfgAldwQAAAABeHEAfgAnc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ACtzcQB+ACwAAABAAgAAQXBwcHQAKWNvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpzZWxlY3RlZFNvcnRzdXEAfgAAAAAAAnVxAH4AAAAAAAJ1cQB+AAAAAAAEcHVxAH4AAAAAAAR+cgAwb3JnLnByaW1lZmFjZXMuY29tcG9uZW50LmFwaS5VSURhdGEkUHJvcGVydHlLZXlzAAAAAAAAAAASAAB4cQB+AAt0AAhyb3dJbmRleHNyABFqYXZhLmxhbmcuSW50ZWdlchLioKT3gYc4AgABSQAFdmFsdWV4cQB+AC3/////fnEAfgBCdAAFc2F2ZWRwcHBzcQB+AAM/QAAAAAAAAHcIAAAAEAAAAAB4c3EAfgADP0AAAAAAAAB3CAAAABAAAAAAeHQAK2NvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpwYXJhbUFkb3B0aW9uTm91cQB+AAAAAAACdXEAfgAAAAAAAXVxAH4AAAAAAARwdXEAfgAAAAAABHEAfgAMc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ABVzcQB+ABtzcQB+AB5zcQB+ACF3PgAqI3t0ZXh0c1snbGVnYWxBY3RTZWFyY2gucGFyYW1BZG9wdGlvbk5vJ119ABBqYXZhLmxhbmcuT2JqZWN0cHB4c3EAfgAjAAAAfgAAAIJ0ABwvcG9ydGFsL2xlZ2FsQWN0U2VhcmNoLnhodG1sdwcABXZhbHVleHEAfgAnc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ACtzcQB+ACwAAAAgAgAAAHBwcHQANGNvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpqX2lkXzZkOnNlbGVjdGVkSXRlbUxpc3R1cQB+AAAAAAACdXEAfgAAAAAAAnVxAH4AAAAAAARwdXEAfgAAAAAABHEAfgBDcQB+AEZxAH4AR3BwcHNxAH4AAz9AAAAAAAAAdwgAAAAQAAAAAHhzcQB+AAM/QAAAAAAAAHcIAAAAEAAAAAB4dAAab2FtLkZBQ0VMRVRfU1RBVEVfSU5TVEFOQ0VzcQB+AA52cgA1b3JnLmFwYWNoZS5teWZhY2VzLnZpZXcuZmFjZWxldHMudGFnLmpzZi5GYWNlbGV0U3RhdGWTbGF69FofrwIAAkwAC2JpbmRpbmdzTWFwdAAPTGphdmEvdXRpbC9NYXA7TAAIc3RhdGVNYXBxAH4AZHhwdXEAfgAAAAAAAXNxAH4AAz9AAAAAAAAMdwgAAAAQAAAACXQAFWI2XzE1MTg0MzUwMTJfZThlMWE1NnNyABFqYXZhLmxhbmcuQm9vbGVhbs0gcoDVnPruAgABWgAFdmFsdWV4cAB0ABU4eV8xNTE4NDM1MDEyX2U4ZTAyODZzcQB+AGkBdAAVOThfMTUxODQzNTAxMl9lOGUwMDRlcQB+AGp0ABU5NF8xNTE4NDM1MDEyX2U4ZTAzYzhxAH4AbHQAFTh6XzE1MTg0MzUwMTJfZThlMDJiOXEAfgBsdAAVOTFfMTUxODQzNTAxMl9lOGUwMzZkcQB+AGx0ABU5Ml8xNTE4NDM1MDEyX2U4ZTAzMDRxAH4AbHQAFTk3XzE1MTg0MzUwMTJfZThlMDM5N3EAfgBqdAAVOTVfMTUxODQzNTAxMl9lOGUwM2UzcQB+AGx4dAAsY29udGVudEZvcm06c2VhcmNoUGFyYW1QYW5lOmpfaWRfM2o6Y2FsZW5kYXJ1cQB+AAAAAAACdXEAfgAAAAAAAXVxAH4AAAAAAARwdXEAfgAAAAAABHEAfgAMc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ABVzcQB+ABZzcQB+ABtzcQB+AB5zcQB+ACF3JQARI3tjYy5hdHRycy5sYWJlbH0AEGphdmEubGFuZy5PYmplY3RwcHhxAH4AJXcHAAV2YWx1ZXhxAH4AJXcEAAAAAXhxAH4AJ3NxAH4ADnEAfgATdXEAfgAAAAAAAnEAfgArc3EAfgAsAAAAQAIAAEFwcHB0ADRjb250ZW50Rm9ybTpzZWFyY2hQYXJhbVBhbmU6al9pZF8yZDpzZWxlY3RlZEl0ZW1MaXN0dXEAfgAAAAAAAnVxAH4AAAAAAAJ1cQB+AAAAAAAEcHVxAH4AAAAAAARxAH4AQ3EAfgBGcQB+AEdwcHBzcQB+AAM/QAAAAAAAAHcIAAAAEAAAAAB4c3EAfgADP0AAAAAAAAB3CAAAABAAAAAAeHQAL2NvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpwYXJhbVJlZ2lzdHJhdGlvbk5vdXEAfgAAAAAAAnVxAH4AAAAAAAF1cQB+AAAAAAAEcHVxAH4AAAAAAARxAH4ADHNxAH4ADnEAfgATdXEAfgAAAAAAAnEAfgAVc3EAfgAbc3EAfgAec3EAfgAhd0IALiN7dGV4dHNbJ2xlZ2FsQWN0U2VhcmNoLnBhcmFtUmVnaXN0cmF0aW9uTm8nXX0AEGphdmEubGFuZy5PYmplY3RwcHhzcQB+ACMAAACGAAAAznEAfgBWdwcABXZhbHVleHEAfgAnc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ACtzcQB+ACwAAAAgAgAAQHBwcHQACWpfaWRfX3ZfMHVxAH4AAAAAAAJwc3EAfgAOdnIAKW9yZy5hcGFjaGUubXlmYWNlcy52aWV3LlZpZXdTY29wZVByb3h5TWFwAAAAAAAAAAAAAAB4cHQACi05NTYyNjkyNzh0ACZjb250ZW50Rm9ybTpzZWFyY2hQYXJhbVBhbmU6cGFyYW1TdGF0ZXVxAH4AAAAAAAJ1cQB+AAAAAAABdXEAfgAAAAAABHB1cQB+AAAAAAAEcQB+AAxzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AFXNxAH4AG3NxAH4AHnNxAH4AIXc5ACUje3RleHRzWydsZWdhbEFjdFNlYXJjaC5wYXJhbVN0YXRlJ119ABBqYXZhLmxhbmcuT2JqZWN0cHB4c3EAfgAjAAAAdAAAARNxAH4AVncHAAV2YWx1ZXhxAH4AJ3NxAH4ADnEAfgATdXEAfgAAAAAAAnEAfgArc3EAfgAsAAAAAAIAAEFwcHB0ACdjb250ZW50Rm9ybTpzZWFyY2hQYXJhbVBhbmU6cGFyYW1Tb3J0Qnl1cQB+AAAAAAACdXEAfgAAAAAAAXVxAH4AAAAAAARwdXEAfgAAAAAABHEAfgAMc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ABVzcQB+ABtzcQB+AB5zcQB+ACF3OgAmI3t0ZXh0c1snbGVnYWxBY3RTZWFyY2gucGFyYW1Tb3J0QnknXX0AEGphdmEubGFuZy5PYmplY3RwcHhzcQB+ACMAAAB2AAABcnEAfgBWdwcABXZhbHVleHEAfgAnc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ACtzcQB+ACwAAAAAAgAAAHBwcHQAKGNvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpzb3J0RHJvcGRvd251cQB+AAAAAAACdXEAfgAAAAAAAXVxAH4AAAAAAARwdXEAfgAAAAAABHEAfgAMc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ABVzcQB+ABtzcQB+AB5zcQB+ACF3QAAsI3t0ZXh0c1snbGVnYWxBY3RTZWFyY2gucGFyYW1Eb2N1bWVudFNvcnQnXX0AEGphdmEubGFuZy5PYmplY3RwcHhzcQB+ACMAAAB9AAAAlHEAfgBWdwcABXZhbHVleHEAfgAnc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ACtzcQB+ACwAAAAAAgAAAXBwcHQAOmNvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpqX2lkXzJkOmRpYWxvZ1NlbGVjdGVkSXRlbUxpc3R1cQB+AAAAAAACdXEAfgAAAAAAAnVxAH4AAAAAAARwdXEAfgAAAAAABHEAfgBDcQB+AEZxAH4AR3BwcHNxAH4AAz9AAAAAAAAAdwgAAAAQAAAAAHhzcQB+AAM/QAAAAAAAAHcIAAAAEAAAAAB4dAAoY29udGVudEZvcm06c2VhcmNoUGFyYW1QYW5lOnBhcmFtQ29udGVudHVxAH4AAAAAAAJ1cQB+AAAAAAABdXEAfgAAAAAABHB1cQB+AAAAAAAEcQB+AAxzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AFXNxAH4AG3NxAH4AHnNxAH4AIXc7ACcje3RleHRzWydsZWdhbEFjdFNlYXJjaC5wYXJhbUNvbnRlbnQnXX0AEGphdmEubGFuZy5PYmplY3RwcHhzcQB+ACMAAAB4AAAAUHEAfgBWdwcABXZhbHVleHEAfgAnc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ACtzcQB+ACwAAABgAgAAAHBwcHQALmNvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpwYXJhbURvY3VtZW50R3JvdXB1cQB+AAAAAAACdXEAfgAAAAAAAXVxAH4AAAAAAARwdXEAfgAAAAAABHEAfgAMc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ABVzcQB+ABtzcQB+AB5zcQB+ACF3QQAtI3t0ZXh0c1snbGVnYWxBY3RTZWFyY2gucGFyYW1Eb2N1bWVudEdyb3VwJ119ABBqYXZhLmxhbmcuT2JqZWN0cHB4c3EAfgAjAAAAhAAAAShxAH4AVncHAAV2YWx1ZXhxAH4AJ3NxAH4ADnEAfgATdXEAfgAAAAAAAnEAfgArc3EAfgAsAAAAAAIAAEFwcHB0ADpjb250ZW50Rm9ybTpzZWFyY2hQYXJhbVBhbmU6al9pZF82ZDpkaWFsb2dTZWxlY3RlZEl0ZW1MaXN0dXEAfgAAAAAAAnVxAH4AAAAAAAJ1cQB+AAAAAAAEcHVxAH4AAAAAAARxAH4AQ3EAfgBGcQB+AEdwcHBzcQB+AAM/QAAAAAAAAHcIAAAAEAAAAAB4c3EAfgADP0AAAAAAAAB3CAAAABAAAAAAeHQALGNvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpqX2lkXzRtOmNhbGVuZGFydXEAfgAAAAAAAnVxAH4AAAAAAAF1cQB+AAAAAAAEcHVxAH4AAAAAAARxAH4ADHNxAH4ADnEAfgATdXEAfgAAAAAAAnEAfgAVc3EAfgAWc3EAfgAbc3EAfgAec3EAfgAhdyUAESN7Y2MuYXR0cnMubGFiZWx9ABBqYXZhLmxhbmcuT2JqZWN0cHB4cQB+ACV3BwAFdmFsdWV4cQB+ACV3BAAAAAF4cQB+ACdzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AK3NxAH4ALAAAAEACAABBcHBwdAA3Y29udGVudEZvcm06c2VhcmNoUGFyYW1QYW5lOnBhcmFtVmFsaWREYXRlRnJvbTpjYWxlbmRhcnVxAH4AAAAAAAJ1cQB+AAAAAAABdXEAfgAAAAAABHB1cQB+AAAAAAAEcQB+AAxzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AFXNxAH4AG3NxAH4AHnNxAH4AIXc9ACkje3RleHRzWydsZWdhbEFjdFNlYXJjaC5wYXJhbVZhbGlkRGF0ZSddfQAQamF2YS5sYW5nLk9iamVjdHBweHNxAH4AIwAAAIkAAADjcQB+AFZ3BwAFdmFsdWV4cQB+ACdzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AK3NxAH4ALAAAAEACAABBcHBwdAAYY29udGVudEZvcm06cmVzdWx0c1RhYmxldXEAfgAAAAAAAnVxAH4AAAAAAAJ1cQB+AAAAAAAEcHVxAH4AAAAAAARxAH4AQ3EAfgBGcQB+AEdwcHBzcQB+AAM/QAAAAAAAAHcIAAAAEAAAAAB4c3EAfgADP0AAAAAAAAB3CAAAABAAAAAAeHQAKWNvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpwYXJhbUFFU2VhcmNodXEAfgAAAAAAAnVxAH4AAAAAAAF1cQB+AAAAAAAEcHVxAH4AAAAAAARxAH4ADHNxAH4ADnEAfgATdXEAfgAAAAAAAnEAfgAVc3EAfgAbc3EAfgAec3EAfgAhdzwAKCN7dGV4dHNbJ2xlZ2FsQWN0U2VhcmNoLnBhcmFtQUVTZWFyY2gnXX0AEGphdmEubGFuZy5PYmplY3RwcHhzcQB+ACMAAACCAAAArXEAfgBWdwcABXZhbHVleHEAfgAnc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ACtzcQB+ACwAAAAAAgAAAHBwcHQAJmNvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpwYXJhbVRpdGxldXEAfgAAAAAAAnVxAH4AAAAAAAF1cQB+AAAAAAAEcHVxAH4AAAAAAARxAH4ADHNxAH4ADnEAfgATdXEAfgAAAAAAAnEAfgAVc3EAfgAbc3EAfgAec3EAfgAhdzkAJSN7dGV4dHNbJ2xlZ2FsQWN0U2VhcmNoLnBhcmFtVGl0bGUnXX0AEGphdmEubGFuZy5PYmplY3RwcHhzcQB+ACMAAAB0AAAAN3EAfgBWdwcABXZhbHVleHEAfgAnc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ACtzcQB+ACwAAABgAgAAAHBwcHQALGNvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpqX2lkXzNuOmNhbGVuZGFydXEAfgAAAAAAAnVxAH4AAAAAAAF1cQB+AAAAAAAEcHVxAH4AAAAAAARxAH4ADHNxAH4ADnEAfgATdXEAfgAAAAAAAnEAfgAVc3EAfgAWc3EAfgAbc3EAfgAec3EAfgAhdyUAESN7Y2MuYXR0cnMubGFiZWx9ABBqYXZhLmxhbmcuT2JqZWN0cHB4cQB+ACV3BwAFdmFsdWV4cQB+ACV3BAAAAAF4cQB+ACdzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AK3NxAH4ALAAAAEACAABBcHBwdAAsY29udGVudEZvcm06c2VhcmNoUGFyYW1QYW5lOnBhcmFtUHVibGlzaGVkQnl1cQB+AAAAAAACdXEAfgAAAAAAAXVxAH4AAAAAAARwdXEAfgAAAAAABHEAfgAMc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ABVzcQB+ABtzcQB+AB5zcQB+ACF3PwArI3t0ZXh0c1snbGVnYWxBY3RTZWFyY2gucGFyYW1QdWJsaXNoZWRCeSddfQAQamF2YS5sYW5nLk9iamVjdHBweHNxAH4AIwAAAIQAAAD7cQB+AFZ3BwAFdmFsdWV4cQB+ACdzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AK3NxAH4ALAAAAAACAABBcHBwdAAuY29udGVudEZvcm06c2VhcmNoUGFyYW1QYW5lOnBhcmFtUHVibGljYXRpb25Ob3VxAH4AAAAAAAJ1cQB+AAAAAAABdXEAfgAAAAAABHB1cQB+AAAAAAAEcQB+AAxzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AFXNxAH4AG3NxAH4AHnNxAH4AIXdBAC0je3RleHRzWydsZWdhbEFjdFNlYXJjaC5wYXJhbVB1YmxpY2F0aW9uTm8nXX0AEGphdmEubGFuZy5PYmplY3RwcHhzcQB+ACMAAACJAAABCXEAfgBWdwcABXZhbHVleHEAfgAnc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ACtzcQB+ACwAAAAgAgAAQHBwcHQAMGNvbnRlbnRGb3JtOnNlYXJjaFBhcmFtUGFuZTpwYXJhbVB1YmxpY2F0aW9uWWVhcnVxAH4AAAAAAAJ1cQB+AAAAAAABdXEAfgAAAAAABHB1cQB+AAAAAAAEcQB+AAxzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AFXNxAH4AG3NxAH4AHnNxAH4AIXdDAC8je3RleHRzWydsZWdhbEFjdFNlYXJjaC5wYXJhbVB1YmxpY2F0aW9uWWVhciddfQAQamF2YS5sYW5nLk9iamVjdHBweHNxAH4AIwAAAI0AAAEHcQB+AFZ3BwAFdmFsdWV4cQB+ACdzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AK3NxAH4ALAAAAGACAABAcHBwdAAxY29udGVudEZvcm06c2VhcmNoUGFyYW1QYW5lOnBhcmFtUHVibGljYXRpb25Eb2NOb3VxAH4AAAAAAAJ1cQB+AAAAAAABdXEAfgAAAAAABHB1cQB+AAAAAAAEcQB+AAxzcQB+AA5xAH4AE3VxAH4AAAAAAAJxAH4AFXNxAH4AG3NxAH4AHnNxAH4AIXdEADAje3RleHRzWydsZWdhbEFjdFNlYXJjaC5wYXJhbVB1YmxpY2F0aW9uRG9jTm8nXX0AEGphdmEubGFuZy5PYmplY3RwcHhzcQB+ACMAAACPAAABC3EAfgBWdwcABXZhbHVleHEAfgAnc3EAfgAOcQB+ABN1cQB+AAAAAAACcQB+ACtzcQB+ACwAAAAgAgAAQHBwcHh0ABwvcG9ydGFsL2xlZ2FsQWN0U2VhcmNoLnhodG1s",
        "contentForm:searchParamPane:paramToggleButtonTop_input": "on",
        "contentForm:searchParamPane:paramTitle": f"{id}",
        "contentForm:searchParamPane:titleSearchOptionSelect": "ALL_WORDS",
        "contentForm:searchParamPane:paramContent": "",
        "contentForm:searchParamPane:textSearchOptionSelect": "ALL_WORDS",
        "contentForm:searchParamPane:j_id_2d:kuku_input": "",
        "contentForm:searchParamPane:j_id_2d:kuku_hinput": "",
        "contentForm:searchParamPane:j_id_2d:filter": "",
        "contentForm:searchParamPane:j_id_2d:tree_selection": "",
        "contentForm:searchParamPane:paramAdoptionNo": "",
        "contentForm:searchParamPane:j_id_3j:calendar_input": "",
        "contentForm:searchParamPane:j_id_3n:calendar_input": "",
        "contentForm:searchParamPane:sortDropdown_focus": "",
        "contentForm:searchParamPane:sortDropdown_input": "",
        "contentForm:searchParamPane:paramRegistrationNo": "",
        "contentForm:searchParamPane:j_id_4m:calendar_input": "",
        "contentForm:searchParamPane:j_id_4q:calendar_input": "",
        "contentForm:searchParamPane:paramValidDateFrom:calendar_input": "",
        "contentForm:searchParamPane:j_id_54:calendar_input": "",
        "contentForm:searchParamPane:paramPublishedBy_focus": "",
        "contentForm:searchParamPane:paramPublishedBy_input": "",
        "contentForm:searchParamPane:paramPublicationYear": "",
        "contentForm:searchParamPane:paramPublicationNo": "",
        "contentForm:searchParamPane:paramPublicationDocNo": "",
        "contentForm:searchParamPane:paramState_focus": "",
        "contentForm:searchParamPane:paramState_input": "",
        "contentForm:searchParamPane:paramDocumentGroup_focus": "",
        "contentForm:searchParamPane:paramDocumentGroup_input": "",
        "contentForm:searchParamPane:j_id_6d:kuku_input": "",
        "contentForm:searchParamPane:j_id_6d:kuku_hinput": "",
        "contentForm:searchParamPane:j_id_6d:thesaurusSearchOptionSelect": "includeRelatedTerms",
        "contentForm:searchParamPane:j_id_6d:filter": "",
        "contentForm:searchParamPane:j_id_6d:tree_selection": "",
        "contentForm:searchParamPane:paramSortBy_focus": "",
        "contentForm:searchParamPane:paramSortBy_input": "registrationDate",
        "contentForm:searchParamPane:paramToggleButton_input": "on",
        "contentForm:searchParamPane_active": "0",
        "contentForm:resultsTable_selection": ""
    }

    response = requests.post(url, data=data, cookies={"CookieName": "CookieValue"}, allow_redirects=True)

    def fix_html_data(html):
        # something to do with weird html structure and XML
        # https://groups.google.com/g/beautifulsoup/c/2yMjUYTIaiQ
        regex = re.compile(r'<!\[CDATA\[(.+?)\]\]>', re.DOTALL)
        html = regex.sub(r'X![CDATA[\1]]X', html)
        return html

    html = fix_html_data(response.text)
    soup = BeautifulSoup(html, "html.parser")

    # find the data table <tbody class="ui-datatable-data ui-widget-content" id="contentForm:resultsTable_data">
    table = soup.find('tbody', {"class": 'ui-datatable-data ui-widget-content', 'id': 'contentForm:resultsTable_data'})

    if 'Nerasta teisės aktų' in table.find('td').text:
        # print('Nerasta teisės aktų', '\n')
        yield {
            'id': id,
        }
        return None

    # move over tr, go to 4 td and first a with href.
    for row in table.find_all('tr'):
        cells = row.find_all('td')
        document = cells[3].find('a', href=True)
        date_accepted = cells[5].text
        date_come_into_force = cells[6].text
        # print(cell_content.text, cell_content['href'])

        doc_data = {
            'id': id,
            'doc': document.text,
            'date_accepted': date_accepted,
            'date_come_into_force': date_come_into_force,
            'doc_url': 'https://www.e-tar.lt' + document['href'],
        }

        # now go to the document page and check for its contents
        doc_url = 'https://www.e-tar.lt' + document['href'].replace('legalActEditions', 'legalAct') + '/asr'

        response_ = requests.get(doc_url)
        document = BeautifulSoup(response_.text)

        # move over the documents priedai
        annexes = document.find_all('li', {'class': 'ui-datalist-item'})
        if len(annexes) > 0:
            for annex in annexes:
                # print(i.text, i.findNext('a')['href'])
                yield doc_data | {
                    'doc_attachment': annex.text,
                    'doc_attachment_url': 'https://www.e-tar.lt' + annex.findNext('a')['href']
                }
        else:
            yield doc_data


def get_project_ids():
    # https://stackoverflow.com/a/60671292
    import ssl
    ssl._create_default_https_context = ssl._create_unverified_context
    url = 'https://www.esinvesticijos.lt/uploads/documents/images/Dokumentai/Kvietim%C5%B3%20planas%202023-05-11%20(2).xlsx'
    df = pd.read_excel(url, header=1)
    return df['Pažangos priemonės numeris'].unique().tolist()


ids = get_project_ids()

# filter out unique and split if multiple ids in single string
ids_ = set()
for id_ in ids:
    if not pd.isna(id_):
        id_ = id_.strip()
        if ' / ' in id_:
            for id__ in id_.split(' / '):
                ids_.add(id__)
        else:
            ids_.add(id_)

datas = []
for id_ in ids_:
    print('\n', '# ------ ', id_, ' ------ #')
    for data in get_data(id_):
        print(data)
        datas.append(data)

# save to excel
df = pd.DataFrame(datas)
df.to_excel('output.xlsx', index=False)

# ----- get PFSA's ----- #


def url_to_doc_to_text(url):
    response = requests.get(url)
    file_obj = io.BytesIO(response.content)
    try:
        document = docx.Document(file_obj)
    except ValueError as e:
        print(e)
        return None

    document_text = ''
    for paragraph in document.paragraphs:
        document_text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                document_text += cell.text + "\t"
            document_text += "\n"

    return document_text


for index, row in df.iterrows():
    url = row['doc_attachment_url']
    if not pd.isna(url):
        try:
            text = url_to_doc_to_text(url)
            df.at[index, 'PFSA'] = text
        except:
            print('error')

df.to_parquet('output.parquet', index=False)
